#!/usr/bin/env python3
"""
Generate a Paradigm form term sheet (.docx) by filling in the template.

Usage:
    python3 generate.py '<JSON parameters>'
    python3 generate.py --file params.json

Output: Term Sheet - {Company} Series {X}.docx
"""

import json
import os
import re
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


DEFAULTS = {
    "no_shop_days": 45,
    "counsel_fee_cap": 75000,
    "debt_limit_m": 1,
    "board_seat": True,
    "observer_seat": True,
    "is_crypto": False,
    "option_pool_percent": 15,
    "token_floor_percent": 50,
}

REQUIRED = ["company_name", "series", "investment_amount", "post_money_valuation"]

TEMPLATE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = TEMPLATE_DIR / "template.docx"


def fmt_m(n):
    """Format number as $XM (millions). E.g. 10000000 -> '10'."""
    m = n / 1_000_000
    if m == int(m):
        return str(int(m))
    return f"{m:.1f}"


def replace_in_runs(paragraph, old, new):
    """Replace text across runs in a paragraph, preserving formatting."""
    # Join all run texts
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False

    full = full.replace(old, new)

    # Redistribute text across runs, keeping first run's formatting
    if paragraph.runs:
        paragraph.runs[0].text = full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def replace_in_cell(cell, old, new):
    """Replace text in all paragraphs of a cell."""
    found = False
    for para in cell.paragraphs:
        if replace_in_runs(para, old, new):
            found = True
    return found


def remove_bracketed_clause(cell, clause_start):
    """Remove a bracketed optional clause like [One director...] from a cell."""
    for para in cell.paragraphs:
        full = "".join(r.text for r in para.runs)
        # Find the bracket that starts with this text
        idx = full.find("[" + clause_start)
        if idx == -1:
            continue
        # Find matching closing bracket
        depth = 0
        end = idx
        for i in range(idx, len(full)):
            if full[i] == "[":
                depth += 1
            elif full[i] == "]":
                depth -= 1
                if depth == 0:
                    end = i
                    break
        # Remove the clause (and any trailing space)
        removed = full[:idx] + full[end + 1:]
        removed = removed.replace("  ", " ").strip()
        if para.runs:
            para.runs[0].text = removed
            for r in para.runs[1:]:
                r.text = ""
        return True
    return False


def clear_cell(cell):
    """Clear all text from a cell while preserving structure."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""


def remove_row(table, row_idx):
    """Remove a row from a table."""
    row = table.rows[row_idx]
    tr = row._tr
    tr.getparent().remove(tr)


def generate(params):
    """Generate the term sheet from template."""
    p = {**DEFAULTS, **params}

    missing = [f for f in REQUIRED if f not in p or p[f] is None]
    if missing:
        print(f"ERROR: Missing required fields: {', '.join(missing)}", file=sys.stderr)
        sys.exit(1)

    company = p["company_name"]
    series = p["series"]
    amount = p["investment_amount"]
    post_money = p["post_money_valuation"]
    pre_money = post_money - amount
    option_pool = p.get("option_pool_percent", 15)
    ownership_pct = (amount / post_money) * 100
    board_seat = p["board_seat"]
    observer_seat = p["observer_seat"]
    is_crypto = p["is_crypto"]
    token_floor = p.get("token_floor_percent", 50)
    counsel_cap = p.get("counsel_fee_cap", 75000)
    no_shop = p.get("no_shop_days", 45)

    # Load template
    if not TEMPLATE_PATH.exists():
        print(f"ERROR: Template not found at {TEMPLATE_PATH}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(TEMPLATE_PATH))

    # === HEADER PARAGRAPHS ===
    # P[0]: "[COMPANY] SERIES [ _] PREFERRED STOCK FINANCING"
    for para in doc.paragraphs:
        replace_in_runs(para, "[COMPANY]", company.upper())
        replace_in_runs(para, "COMPANY]", company.upper() + "]")
        replace_in_runs(para, "[ _]", series)
        replace_in_runs(para, "[_]", series)

    # === SIGNATURE TABLE (Table 1) ===
    sig_table = doc.tables[1]
    for row in sig_table.rows:
        for cell in row.cells:
            replace_in_cell(cell, "[COMPANY]", company.upper())

    # === MAIN TERMS TABLE (Table 0) ===
    table = doc.tables[0]

    # --- Row 0: Investment & Post-Money Valuation ---
    cell = table.rows[0].cells[1]
    # Replace investment amount: $[__]M (first occurrence)
    replace_in_cell(cell, "$[__]M at a $[__]M", f"${fmt_m(amount)}M at a ${fmt_m(post_money)}M")
    # Option pool percentage
    replace_in_cell(cell, "[__] % of the post-money", f"{option_pool}% of the post-money")
    # Ownership percentage
    replace_in_cell(cell, "[__] % of the fully diluted", f"{ownership_pct:.1f}% of the fully diluted")
    # Company name
    replace_in_cell(cell, "[COMPANY]", company)
    replace_in_cell(cell, "COMPANY]", company + "]")
    # Co-investor clause - remove the optional bracket
    remove_bracketed_clause(cell, "Other investors")

    # Also try individual [__] replacements that might not have been caught
    # Pre-money valuation: $[__]M for pre-money
    full_text = cell.text
    if "$[__]" in full_text:
        replace_in_cell(cell, "$[__]", f"${fmt_m(pre_money)}")

    # --- Row 1: Securities ---
    cell = table.rows[1].cells[1]
    replace_in_cell(cell, "[__]", series)

    # --- Row 2: Board and Voting Rights ---
    cell = table.rows[2].cells[1]
    replace_in_cell(cell, "[__]", series)

    if not board_seat and not observer_seat:
        # Remove both board seat and observer clauses
        remove_bracketed_clause(cell, "One director")
        remove_bracketed_clause(cell, "In addition")
    elif board_seat and not observer_seat:
        # Keep board seat, remove observer
        remove_bracketed_clause(cell, "In addition")
    elif not board_seat and observer_seat:
        # Remove board seat, keep observer
        remove_bracketed_clause(cell, "One director")
    # else: both True, keep both clauses as-is (just remove the outer brackets)
    if board_seat:
        # Remove outer brackets from board clause but keep text
        for para in cell.paragraphs:
            full = "".join(r.text for r in para.runs)
            # Replace "[One director..." with "One director..."
            full = re.sub(
                r'\[One director to be elected by the Series .+? Preferred Stock and designated by Paradigm\.\]',
                lambda m: m.group(0)[1:-1],  # strip outer [ ]
                full
            )
            if observer_seat:
                full = re.sub(
                    r'\[In addition, Company shall invite.*?directors;\]',
                    lambda m: m.group(0)[1:-1],
                    full,
                    flags=re.DOTALL
                )
            if para.runs:
                para.runs[0].text = full
                for r in para.runs[1:]:
                    r.text = ""

    # --- Row 3: Protective Provisions ---
    cell = table.rows[3].cells[1]
    debt_limit = p.get("debt_limit_m", 1)
    replace_in_cell(cell, "[1-10M]", str(debt_limit))

    # --- Row 4: Other Rights ---
    # No placeholders to fill

    # --- Row 5: Token Rights ---
    if not is_crypto:
        # Remove the entire Token Rights row
        remove_row(table, 5)
    else:
        cell = table.rows[5].cells[1]
        replace_in_cell(cell, "[Token Floor usually 50%]", str(token_floor))

    # --- Row 6 (or 5 if token row removed): Vesting ---
    # No placeholders to fill (standard language)

    # --- Row 7 (or 6): Documentation; Legal Fees ---
    # Counsel fee cap is already $75,000 in template
    # Only replace if different from default
    row_idx = 7 if is_crypto else 6
    if row_idx < len(table.rows):
        cell = table.rows[row_idx].cells[1]
        if counsel_cap != 75000:
            replace_in_cell(cell, "$75,000", f"${counsel_cap:,.0f}")

    # --- Row 8 (or 7): No-Shop; Confidentiality ---
    row_idx = 8 if is_crypto else 7
    if row_idx < len(table.rows):
        cell = table.rows[row_idx].cells[1]
        if no_shop != 45:
            replace_in_cell(cell, "45 days", f"{no_shop} days")

    # === REMOVE DRI NOTES (paragraphs starting with "Paradigm DRI to confirm") ===
    paras_to_remove = []
    dri_found = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Paradigm DRI to confirm"):
            dri_found = True
        if dri_found and text:
            paras_to_remove.append(para)
    for para in paras_to_remove:
        parent = para._element.getparent()
        parent.remove(para._element)

    # === STRIP ALL HIGHLIGHTING ===
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.highlight_color = None
    for table_obj in doc.tables:
        for row in table_obj.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.highlight_color = None

    # === SAVE ===
    output_dir = p.get("output_dir", os.getcwd())
    filename = f"Term Sheet - {company} Series {series}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: generate.py '<JSON>' or generate.py --file params.json", file=sys.stderr)
        sys.exit(1)

    if sys.argv[1] == "--file":
        with open(sys.argv[2]) as f:
            params = json.load(f)
    else:
        params = json.loads(sys.argv[1])

    generate(params)


if __name__ == "__main__":
    main()
